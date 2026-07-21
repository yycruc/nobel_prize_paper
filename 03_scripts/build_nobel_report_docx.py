# -*- coding: utf-8 -*-
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "02_full_collection" / "08_report"
CHART_DIR = REPORT_DIR / "charts"
OUT_DOCX = REPORT_DIR / "诺奖关键论文分析报告.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        set_cell_shading(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
    doc.add_paragraph()


def add_para(doc: Document, text: str, style: str | None = None, bold_lead: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_lead and text.startswith(bold_lead):
        r1 = p.add_run(bold_lead)
        r1.bold = True
        r1.font.name = "微软雅黑"
        r1._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        r2 = p.add_run(text[len(bold_lead):])
        r2.font.name = "微软雅黑"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    else:
        run = p.add_run(text)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return p


def add_caption(doc: Document, text: str) -> None:
    p = add_para(doc, text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(90, 90, 90)


def add_chart(doc: Document, filename: str, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(CHART_DIR / filename), width=Cm(15.5))
    add_caption(doc, caption)


def set_doc_style(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(5)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "微软雅黑"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.color.rgb = RGBColor(31, 78, 121)
    doc.styles["Heading 1"].font.size = Pt(15)
    doc.styles["Heading 2"].font.size = Pt(12.5)
    doc.styles["Heading 3"].font.size = Pt(11.5)


def build() -> None:
    doc = Document()
    set_doc_style(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("从诺奖关键论文看科学奖项产出规律与我国启示")
    r.bold = True
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("基于诺奖官网线索、Li 等公开数据、OpenAlex 匹配及代表性期刊发文数据")
    sr.font.name = "微软雅黑"
    sr._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    sr.font.size = Pt(10.5)

    add_para(doc, "说明：本文按原 Word 提纲梳理前期数据采集与分析结论，并将图表单独整理在 Excel 工作簿中生成后嵌入本文。")

    doc.add_heading("一、研究方法与数据基础", level=1)
    add_para(
        doc,
        "本次分析采用“先官方线索、再外部数据补充、最后 OpenAlex 统一匹配”的保守口径。"
        "自然科学类诺奖基准记录为 662 条（按获奖人-奖项记录计），最终主分析采用 504 条已形成较高可信关键论文匹配的记录，"
        "对应 823 条关键论文行、709 篇去重论文、153 个期刊/来源。其余记录主要进入人工核验或补充队列，避免把纪念性论文、综述、教材、实验报告或后验引用误当成获奖奠基论文。"
    )
    add_table(
        doc,
        ["环节", "数据来源与处理", "主要结果"],
        [
            ["1. 诺奖记录基准", "诺贝尔奖官网自然科学类获奖信息，统一为获奖人-奖项记录。", "1901-2025 年共 662 条自然科学类诺奖记录。"],
            ["2. 官网论文线索", "抓取 Nobel Prize 官方页面、参考文献 PDF、Further reading 等线索，抽取论文、书籍、综述等参考条目。", "官网页面目标 3310 个，成功 2084 个；PDF 链接 1537 个，成功 1535 个；抽取参考文献约 15331 条。"],
            ["3. 公开数据补充", "引入 Li 等整理的诺奖论文数据，作为已有学术数据集补充。", "源记录 874 条，形成候选 863 条，覆盖约 530/662 条诺奖记录。"],
            ["4. OpenAlex 匹配", "用 DOI、题名、年份、期刊和作者等字段匹配 OpenAlex，统一论文、期刊、年份和国家口径。", "候选论文 974 行覆盖 563 条诺奖记录；主分析保留 823 行，覆盖 504 条记录。"],
            ["5. 国家发文趋势", "以诺奖高频期刊 Top77 与 Top20 为观察对象，按 OpenAlex 论文作者国家参与计数计算年度占全球比例。", "Top77 覆盖约 90.0% 的诺奖关键论文行；形成 1980-2025 年六国年度占比数据，Top20 补充到更早历史段。"],
        ],
    )
    add_para(
        doc,
        "口径提醒：国家发文占比使用 OpenAlex authorships.countries 的参与计数，一篇论文若有多个国家参与，会分别计入相关国家一次；分母为该期刊-年份的全球论文量。因此六个国家比例不能简单加总。"
    )

    doc.add_heading("二、从诺奖奠基论文看诺奖规律", level=1)
    doc.add_heading("1. 诺奖成果与论文发表关系", level=2)
    add_para(
        doc,
        "从可匹配记录看，诺奖自然科学成果多数能追溯到一组关键论文，而不是单篇论文。"
        "主分析中的 504 条获奖记录对应 823 条关键论文行，说明不少获奖成果需要由连续论文、同一发现链条或共同方法平台共同支撑。"
        "这也解释了为什么官方材料与外部数据集必须结合使用：官网线索覆盖权威但结构不统一，外部数据结构化程度较高但时间和记录覆盖不完整。"
    )

    doc.add_heading("2. 获奖周期整体拉长", level=2)
    add_para(
        doc,
        "关键论文发表到最终获奖的平均时滞约 17.0 年，中位数约 15 年，最长可达 56 年。"
        "按年代看，早期诺奖从论文到获奖往往不足 10 年；二战后，尤其 1990 年代以后，中位时滞稳定抬升到 20 年左右。"
        "这反映出现代自然科学发现的验证周期、学科共同体确认、仪器平台和后续影响扩散都更长。"
    )
    add_para(
        doc,
        "2020 年代曲线出现急剧缩短，不应解读为长期趋势反转。原因是 2020 年代本身尚未结束，且近期获奖往往包含已经被较快确认的突破；更晚才会获奖的长时滞成果还没有出现在样本中，存在明显右截尾。"
    )
    add_chart(doc, "fig2_lag_by_category.png", "图 1 关键论文发表到诺奖获奖的时滞变化（中位数，按获奖年代）")

    doc.add_heading("3. 关键论文期刊高度集中", level=2)
    add_para(
        doc,
        "153 个期刊/来源中，前 10 位已覆盖约 54% 的关键论文行；扩展到 Top77 期刊后覆盖约 90%。"
        "这说明诺奖成果并不只出现在少数综合期刊，但代表性期刊具有明显集中度，尤其 Nature、Physical Review、Physical Review Letters、Science、JACS、PNAS 等长期占据核心位置。"
    )
    add_table(
        doc,
        ["排名", "期刊", "关键论文行", "覆盖诺奖记录"],
        [
            ["1", "Nature", "99", "83"],
            ["2", "Physical Review", "73", "59"],
            ["3", "Physical Review Letters", "64", "51"],
            ["4", "Science", "42", "34"],
            ["5", "Journal of the American Chemical Society", "41", "28"],
            ["6", "Proceedings of the National Academy of Sciences", "34", "26"],
            ["7", "Journal of Biological Chemistry", "30", "22"],
            ["8", "Cell", "24", "20"],
            ["9", "The Journal of Physiology", "23", "11"],
            ["10", "The Journal of Chemical Physics", "15", "10"],
        ],
    )
    add_chart(doc, "fig3_journal_coverage.png", "图 2 诺奖关键论文期刊集中度：Top77 期刊覆盖约 90% 关键论文行")

    doc.add_heading("4. 国家获奖时间与美国战后领先", level=2)
    add_para(
        doc,
        "PDF 中各国获奖年代数据表明，美国在 1930 年代开始接近并超过欧洲主要国家，二战后迅速拉开差距。"
        "截至 2021 年，科学类诺奖获奖人次中美国约 269 人次，超过 40%；英国、德国、法国、日本分别为约 94、73、34、22 人次。"
        "这一变化与战时科研组织、战后长期稳定研发投入、大学-国家实验室-产业体系联动高度相关。"
    )
    add_chart(doc, "fig1_awards_by_decade.png", "图 3 主要国家诺贝尔自然科学奖获奖数年代分布（PDF 数据，至 2021 年）")

    doc.add_heading("三、我国目前发展情况分析", level=1)
    add_para(
        doc,
        "从代表性诺奖论文发表期刊看，中国论文产出占比已经快速上升。"
        "在 Top77 诺奖高频期刊中，中国占全球比例从 2000 年前后的低个位数上升到 2020-2025 年约 29%，2025 年约 30%；美国同期约 32%-34%，仍处第一梯队但差距明显缩小。"
        "在更严格的 Top20 期刊口径下，中国 2025 年约 21%，上升同样明显，但低于 Top77，说明越聚焦顶级核心期刊，追赶仍更困难。"
    )
    add_chart(doc, "fig4_top77_country_share.png", "图 4 六国在 Top77 诺奖高频期刊中的年度发文占全球比例（1980-2025）")
    add_chart(doc, "fig5_china_top77_top20.png", "图 5 中国在 Top77 与 Top20 诺奖高频期刊中的发文占比（2000-2025）")
    add_para(
        doc,
        "如果参照诺奖 15-20 年甚至更长的平均确认周期，中国 2000 年以后在代表性期刊中的快速增长，更可能在 2030 年代以后逐步体现为奖项候选基础。"
        "但这一判断不能机械外推：诺奖奖励的是被长期验证的原创发现、方法或机制突破，而不是论文数量本身。"
    )
    add_para(
        doc,
        "AI 可能缩短发现和验证的部分环节，但也会提高全球研究产出的整体速度。对我国而言，关键不是单纯提高论文数量，而是能否形成原创问题、关键仪器平台、可复用数据资源和具有国际解释力的学术叙事。"
    )

    doc.add_heading("四、有关建议", level=1)
    recommendations = [
        "围绕原创科学问题进行长期布局。诺奖时滞显示重大成果往往需要 15-20 年以上持续验证，评价和资助机制应允许长周期积累。",
        "把代表性期刊发文作为观察窗口，而不是唯一目标。Top77/Top20 能反映国际可见度，但真正关键是原创性、可重复性和后续影响。",
        "强化国家级科研平台、开放数据和仪器方法能力。美国战后经验表明，大科学组织能力和稳定投入是持续产生突破的重要基础。",
        "建立诺奖潜在成果跟踪库。对中国学者在高频诺奖期刊、重大仪器/方法、基础机制发现中的论文进行持续人工核验，区分高被引论文、综述论文和真正原创奠基论文。",
        "提高国际合作和学术叙事能力。国家计数口径显示国际合作论文会同时计入多国，未来应同时关注贡献位置、通讯作者、核心方法归属和后续引用链条。"
    ]
    for item in recommendations:
        add_para(doc, "• " + item)

    doc.add_heading("附：数据口径与局限", level=1)
    add_para(
        doc,
        "本报告的诺奖关键论文主分析采用保守口径，优先纳入 Nobel 官网、Li 等数据集和 OpenAlex 可较高可信匹配的论文。"
        "因此 504 条主分析记录不是全部 662 条自然科学诺奖记录，也不是全部 974 条候选论文；差异主要来自人工核验队列、历史非期刊文献、书籍/章节、综述性材料和缺少可靠题名/年份/期刊匹配的记录。"
    )
    add_para(
        doc,
        "PDF 中各国获奖时间数据截至 2021 年，与本项目 1901-2025 年诺奖记录口径存在时间差；本文只用其作为国家获奖年代趋势的补充依据。"
        "代表性期刊发文占比来自 OpenAlex 年度论文数据，历史早期数据覆盖度弱于现代数据，跨国合作论文采用国家参与计数。"
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
